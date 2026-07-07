#!/usr/bin/env python3
"""Update RepoSentinel AI Product Plan docx with monetization, design, and roadmap."""

from pathlib import Path
import shutil

from docx import Document

SRC = Path("/data/poc/reposentinelAI/RepoSentinel_AI_Product_Plan.docx")
BACKUP = Path("/data/poc/reposentinelAI/RepoSentinel_AI_Product_Plan_v0.1_backup.docx")


def stage_heading(doc, new_content, text, level=1):
    h = doc.add_heading(text, level=level)
    h._element.getparent().remove(h._element)
    new_content.append(h._element)


def stage_para(doc, new_content, text):
    p = doc.add_paragraph(text)
    p._element.getparent().remove(p._element)
    new_content.append(p._element)


def stage_bullet(doc, new_content, text, bold_prefix=None):
    p = doc.add_paragraph()
    p._element.getparent().remove(p._element)
    p.add_run("• ")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    new_content.append(p._element)


def main():
    shutil.copy2(SRC, BACKUP)
    doc = Document(str(SRC))

    for p in doc.paragraphs[:6]:
        if "v0.1" in p.text:
            for run in p.runs:
                run.text = run.text.replace(
                    "v0.1 Draft", "v0.2 — Monetization & Design Update"
                )

    body = doc.element.body
    roadmap_el = None
    for child in body:
        if child.tag.endswith("}p"):
            texts = "".join(t.text or "" for t in child.iter() if t.tag.endswith("}t"))
            if texts.strip().startswith("13. Phased Roadmap"):
                roadmap_el = child
                break

    if roadmap_el is None:
        raise SystemExit("roadmap element not found")

    new_content = []

    stage_heading(doc, new_content, "13. Monetization & Premium Capabilities", 1)
    stage_para(
        doc,
        new_content,
        'The wedge product (unified scanning, flags, knowledge base) is useful. The capabilities '
        "below move RepoSentinel AI from useful tool to thing people pay for — grouped by what "
        "each addition is selling to the buyer.",
    )

    capabilities = [
        (
            "1. Auto-remediation, not just detection",
            "Flagging a problem is a cost center; fixing it is a budget line. Add a Remediation "
            "Agent that opens a draft PR with the fix (dependency bump, secret rotation "
            "instructions, IaC config correction) alongside every actionable flag. This is the "
            "single biggest lever on perceived value — competitors like Apiiro and Aikido lead "
            'marketing with "auto-fix" because "we found 400 issues" sells nothing but "we closed '
            '60 of them automatically" sells itself.',
        ),
        (
            "2. A single posture score, trending over time",
            "CISOs need something for the board deck. A composite posture score (credit-score "
            "style) per repo, team, and org — with trend lines and peer benchmarking — turns the "
            "dashboard from an engineering tool into an executive reporting tool. Board-reportable "
            "metrics get sponsored; engineering tools get expensed.",
        ),
        (
            "3. Compliance mapping out of the box",
            "Every enterprise buyer is really buying an easier SOC 2 / ISO 27001 / PCI-DSS audit. "
            "Pre-built control mappings with exportable, auditor-ready evidence packages (not "
            "dashboard screenshots) is what turns a security tool into something Legal and "
            "Compliance will co-sign the purchase order for.",
        ),
        (
            "4. M&A / vendor-risk due diligence mode",
            "A dedicated engagement mode that lets a consultancy point the tool at a target "
            "company's repos (with access) and produce a codebase health + security debt report "
            "in hours, not weeks. This is a different sales motion — a single high-value "
            "engagement, not a subscription — and maps directly onto the Tool Hub / consultancy "
            "buyer already in scope.",
        ),
        (
            "5. New-hire / onboarding mode for the knowledge base",
            'Reframe the knowledge base front door around "ask this instead of pinging a senior '
            'engineer." Sellable to engineering managers on ramp-time reduction alone, completely '
            "separate from the security buyer — widening who inside a company can champion the "
            "purchase.",
        ),
        (
            "6. Multi-tenant / white-label for MSPs and consultancies",
            "For technical-services firms managing multiple clients' repos: one pane of glass per "
            "client with strict tenant isolation and consultancy rebrand capability in front of "
            "their own clients. This turns one sale into N (every client engagement) instead of "
            "one.",
        ),
        (
            "7. Native presence in marketplaces buyers already browse",
            "GitHub Marketplace, GitLab integrations page, Atlassian Marketplace. This is a "
            "distribution channel, not a feature — but it is often the actual reason a "
            "competitor gets evaluated and RepoSentinel does not.",
        ),
        (
            "8. A free / PLG tier",
            "Aikido and Snyk both grew through a generous free tier for open-source or small "
            "repos. Given the token-efficiency architecture, Tier-0 scanning can be given away "
            "cheaply as lead-gen, with paid tiers for correlation, knowledge base, auto-remediation, "
            "and compliance exports.",
        ),
        (
            "9. AI-agent and MCP governance",
            'A 2026-specific opportunity: as companies wire AI coding agents and MCP servers into '
            "pipelines, almost nobody has visibility into what those agents touched or exposed. A "
            'module that tracks AI-agent-authored changes and MCP server permissions — an "AI Bill '
            'of Materials for AI agents" — is timely and differentiated.',
        ),
        (
            "10. Deployment flexibility for regulated buyers",
            "A VPC / self-hosted deployment option (reusing the self-hosted vLLM pattern from AIL) "
            "unlocks banks, governments, and crypto-native clients who will not send code to a "
            "SaaS multi-tenant service — a segment the existing fintech/crypto client base already "
            "sits in.",
        ),
    ]

    for title, body_text in capabilities:
        stage_bullet(doc, new_content, body_text, bold_prefix=f"{title}. ")

    stage_heading(doc, new_content, "13.1 Build-Effort vs. Sales-Impact Ranking", 2)
    stage_para(
        doc,
        new_content,
        "Each capability is scored on sales impact and build effort. Priority = high impact "
        "relative to effort, sequenced to compound distribution and revenue.",
    )

    ranking_rows = [
        ("Auto-remediation (draft PR)", "Very High", "Medium-High", "V1 — highest ROI after wedge"),
        ("Posture score + trending", "Very High", "Medium", "V1 — unlocks executive sponsorship"),
        ("Free / PLG tier", "High", "Low", "MVP launch — lead-gen from day one"),
        ("M&A due diligence mode", "Very High (consultancy)", "Medium", "V1 — premium engagement SKU"),
        ("New-hire onboarding mode", "Medium-High", "Low-Medium", "V1 — second buyer persona"),
        ("Marketplace listings", "High (distribution)", "Low-Medium", "Parallel from V1"),
        ("Compliance evidence packages", "Very High", "High", "V2 — Legal/Compliance co-sign"),
        ("Multi-tenant white-label", "High", "High", "V2 — MSP/consultancy scale motion"),
        ("VPC / self-hosted deployment", "High (regulated)", "High", "V2 — banking/crypto segment"),
        ("AI-agent & MCP governance", "Medium-High (differentiated)", "Medium-High", "V2 — 2026 play"),
    ]

    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Capability"
    hdr[1].text = "Sales Impact"
    hdr[2].text = "Build Effort"
    hdr[3].text = "Roadmap Slot"
    table_el = table._element
    table_el.getparent().remove(table_el)
    new_content.append(table_el)
    for cap, impact, effort, slot in ranking_rows:
        row = table.add_row().cells
        row[0].text = cap
        row[1].text = impact
        row[2].text = effort
        row[3].text = slot

    stage_heading(
        doc, new_content, "14. Product Design System — Silicon Valley Clean Aesthetic", 1
    )
    stage_para(
        doc,
        new_content,
        "RepoSentinel AI should feel like a modern infrastructure product (Linear, Vercel, "
        "Datadog) — not a legacy security console. Design is a trust signal for CISOs and platform "
        'teams evaluating whether this is "another noisy scanner" or a product engineers will use.',
    )

    stage_heading(doc, new_content, "14.1 Design Principles", 2)
    principles = [
        ("Clarity over density", "One primary action per screen. Progressive disclosure for drill-downs."),
        ("Data with narrative", "Every chart answers: Are we getting safer? What changed this week?"),
        ("Trust through traceability", "Every score and AI explanation links to file, rule, and commit."),
        ("Calm by default, urgent when needed", "Neutral palette; severity color reserved for actionable risk."),
        ("Speed as a feature", "Sub-200ms navigation, skeleton loaders, optimistic UI on flag assignment."),
    ]
    for title, body_text in principles:
        stage_bullet(doc, new_content, body_text, bold_prefix=f"{title}. ")

    stage_heading(doc, new_content, "14.2 Visual Language", 2)
    for v in [
        "Typography: Inter or Geist for UI; JetBrains Mono for code citations.",
        "Color: Near-white or deep slate backgrounds; single accent for primary actions; semantic severity scale only.",
        "Layout: 12-column grid, max 1280px, collapsible sidebar (Posture, Flags, Architecture, Knowledge, Compliance).",
        "Components: shadcn/ui + Tailwind on Next.js; Tremor for trends; React Flow for architecture graph.",
        "Motion: 150–200ms ease-out transitions; no gratuitous animation.",
        'Empty states: Actionable ("Connect your first repo") — never a blank table.',
    ]:
        stage_bullet(doc, new_content, v)

    stage_heading(doc, new_content, "14.3 Key Screens (MVP → V2)", 2)
    screens = [
        ("Executive Posture Dashboard", "Hero score, 90-day trend, peer benchmark, auto-remediation stats."),
        ("Risk Register", "Filterable table, SLA timers, one-click Create fix PR on eligible flags."),
        ("Flag Detail", "Findings left, agent explanation right, full audit trail footer."),
        ("Architecture Map", "Interactive service graph with drift highlights."),
        ("Knowledge Copilot", "Chat with inline citations; onboarding mode toggle for new hires."),
        ("Compliance Center", "Control framework picker, export evidence package (PDF + JSON)."),
        ("Due Diligence Report", "White-label M&A report — executive summary + technical appendix."),
    ]
    for title, body_text in screens:
        stage_bullet(doc, new_content, body_text, bold_prefix=f"{title}. ")

    stage_heading(doc, new_content, "14.4 White-Label & Multi-Tenant UX", 2)
    stage_para(
        doc,
        new_content,
        "Consultancies get tenant switcher, per-client branding (logo, accent, custom domain), and "
        "client-facing report templates that hide RepoSentinel branding when white-label is enabled.",
    )

    idx = list(body).index(roadmap_el)
    for i, el in enumerate(new_content):
        body.insert(idx + i, el)

    renumber_map = {
        "13. Phased Roadmap": "15. Phased Roadmap",
        "14. Risks & Mitigations": "16. Risks & Mitigations",
        "15. Strategic Fit with Tool Hub & Build Readiness Engine": (
            "17. Strategic Fit with Tool Hub & Build Readiness Engine"
        ),
    }
    for p in doc.paragraphs:
        for old, new in renumber_map.items():
            if p.text.strip().startswith(old):
                for run in p.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)

    roadmap_updates = {
        "MVP (6–8 weeks) — prove the wedge": [
            "MVP (6–8 weeks) — prove the wedge",
            "• GitHub connector, secrets + SCA scanning, clean UI shell (Posture, Risk Register, Flag Detail).",
            "• Triage + Explainer Agents (Tier 0/1 only).",
            "• Free / PLG tier: Tier-0 scanning for small repos — lead-gen from launch.",
            "• Manual KB seeding on 1–2 live Tool Hub pursuits.",
            "• Design system: typography, tokens, shadcn/ui, dark mode.",
        ],
        "V1 (next quarter) — multi-VCS + full deterministic core": [
            "V1 (next quarter) — monetization wedge",
            "• GitLab + Bitbucket; full scanner suite (SAST, IaC, license, drift).",
            "• Auto-remediation Agent: draft PR with fix alongside flags.",
            "• Composite posture score with 90-day trend.",
            "• New-hire onboarding mode; M&A due diligence engagement mode.",
            "• Correlation + Knowledge-Curator Agents; cost dashboard.",
            "• GitHub Marketplace listing.",
        ],
        "V2 — platform tier": [
            "V2 — platform tier & enterprise",
            "• Compliance Center with auditor-ready evidence export (SOC 2, ISO 27001, PCI-DSS).",
            "• Multi-tenant white-label for MSPs/consultancies.",
            "• AI-agent & MCP governance module.",
            "• VPC / self-hosted for regulated buyers.",
            "• Tier 2 escalation; productized pricing; peer benchmarking.",
            "• GitLab + Atlassian Marketplace listings.",
        ],
    }

    for p in doc.paragraphs:
        t = p.text.strip()
        for heading, lines in roadmap_updates.items():
            if t.startswith(heading.split("\n")[0].split("•")[0].strip()):
                p.clear()
                p.add_run(lines[0]).bold = True
                for line in lines[1:]:
                    p.add_run("\n" + line)
                break

    for p in doc.paragraphs:
        if p.text.strip().startswith("16. Risks & Mitigations"):
            risks_el = p._element
            body = doc.element.body
            pricing_els = []
            h = doc.add_heading("15.1 Pricing & Packaging (Draft)", 2)
            h._element.getparent().remove(h._element)
            pricing_els.append(h._element)
            tiers = [
                ("Tier 0 — Free (PLG)", "Secrets + SCA on up to 5 repos. No auto-fix or compliance export."),
                ("Tier 1 — Team", "Posture score, KB, flag workflow, Slack/Jira. Per-seat pricing."),
                ("Tier 2 — Business", "Auto-remediation, compliance mapping, M&A mode, onboarding mode."),
                ("Tier 3 — Enterprise / MSP", "White-label, multi-tenant, VPC/self-hosted, AI-agent governance."),
            ]
            for title, tier_body in tiers:
                bp = doc.add_paragraph()
                bp._element.getparent().remove(bp._element)
                bp.add_run("• ")
                r = bp.add_run(title + ": ")
                r.bold = True
                bp.add_run(tier_body)
                pricing_els.append(bp._element)
            ridx = list(body).index(risks_el)
            for i, el in enumerate(pricing_els):
                body.insert(ridx + i, el)
            break

    doc.save(str(SRC))
    print(f"Updated: {SRC}")
    print(f"Backup: {BACKUP}")


if __name__ == "__main__":
    main()
